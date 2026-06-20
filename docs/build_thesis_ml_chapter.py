from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("thesis-ml-chapter.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"


def set_font(run, size: float, color: RGBColor = INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_mar = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                element = tc_mar.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    tc_mar.append(element)
                element.set(qn("w:w"), "80" if side in {"top", "bottom"} else "120")
                element.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers, strict=True):
        shade(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_font(run, 8.5, INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            set_font(run, 8.5, INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    set_font(run, 9, MUTED, italic=True)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    field_char = OxmlElement("w:fldChar")
    field_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_char, instr, separate, end])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for level, size, color, before, after in ((1, 16, BLUE, 18, 10), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("F1 Prediction Platform | Machine Learning Chapter")
    set_font(header_run, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Page ")
    set_font(footer_run, 8.5, MUTED)
    add_page_number(footer)


def build() -> None:
    doc = Document()
    configure(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(8)
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("THESIS CHAPTER DRAFT"), 10.5, DARK_BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("Machine Learning Experiment Design and Results"), 26, INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("Formula 1 Driver-Level Prediction Across Pre- and Post-Qualifying Information Contexts"), 13.5, MUTED, italic=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("Final experiment: thesis-final-2025-holdout-20260620-r3 | Held-out season: 2025 | Seed: 42"), 9.5, MUTED)

    add_heading(doc, "1. Research Objective and Research Questions", 1)
    add_body(doc, "The objective of this study is to evaluate whether conventional linear models and tree-based machine-learning models can provide reliable driver-level Formula 1 predictions when restricted to information available before a race. The contribution is not a claim that one algorithm universally predicts Formula 1 outcomes. Rather, the study establishes a reproducible and leakage-aware comparison of common candidate families across four existing tasks and two operational information contexts: pre-qualifying and post-qualifying.")
    add_body(doc, "Each observation represents a driver-race row. Predictions are made either before qualifying or after qualifying but before the race. The methodological emphasis is chronological generalisation: model selection is separated from a later, completed-season evaluation, and the final held-out season is not used to choose an algorithm, feature subset, threshold, or champion.")
    add_bullets(doc, [
        "RQ1: Across the four prediction tasks, which candidate model families improve on no-skill baselines under chronological validation?",
        "RQ2: To what extent does the post-qualifying information set improve predictive performance relative to the pre-qualifying information set?",
        "RQ3: Do validation-selected models retain useful performance and probability calibration on an unseen completed season?",
        "RQ4: For Top 10 and podium classification, do the estimated probabilities show adequate discrimination and calibration for the analytical application?",
    ])

    add_heading(doc, "2. Dataset and Temporal Data Partitioning", 1)
    add_body(doc, "The final experiment used only seasons that passed the repository's completed-season checks. The data audit identified 2021, 2023, 2024, and 2025 as fully scheduled with complete race and qualifying coverage. Season 2022 was excluded because race-result rows were absent despite qualifying coverage. Season 2026 was excluded because it contained future scheduled races and partial results. Consequently, no incomplete 2026 data influenced training, validation, threshold selection, or final evaluation.")
    add_table(doc, ["Role", "Seasons", "Purpose"], [
        ["Development pool", "2021, 2023, 2024", "Completed seasons with eligible features and labels."],
        ["Chronological validation", "Train 2021 and 2023; validate 2024", "The sole rolling-origin validation fold; all training races precede validation races."],
        ["Final holdout", "2025", "Latest completed season; not used for selection or threshold fitting."],
    ], [1800, 2700, 4860])
    add_caption(doc, "Suggested Table 1. Temporal partitioning used in the final experiment. Source: experiment manifest and final experiment results artifact.")
    add_body(doc, "The missing 2022 results limited the available development evidence to one outer validation fold. The final run therefore used the earliest feasible expanding split with two training seasons, followed by validation on 2024. Models were then refit on 2021, 2023, and 2024 and evaluated once on 2025. This partition is leakage-safe, but it gives less stable cross-season evidence than the originally desired design with at least three training seasons in every development fold.")
    add_body(doc, "The feature snapshot contained 420, 440, 479, and 479 pre-qualifying rows in 2021, 2023, 2024, and 2025, respectively; the corresponding post-qualifying counts were 439, 440, 479, and 479. Missing target labels were not imputed. Regression tasks used only their explicitly eligible rows, while the classification labels were derived from available race-result rows.")

    add_heading(doc, "3. Feature Engineering and Leakage Prevention", 1)
    add_body(doc, "The pre-qualifying information set included historical pace, driver recent form, team recent form, circuit-history average finishing position, circuit-history DNF rate, and recent DNF rate. Historical calculations were constrained to races earlier than the target race. Driver form used up to five prior races, team form used the team's three prior races, recent DNF rate used up to ten prior races, and pace used clean-lap information from up to three earlier races. Circuit-history features used earlier appearances at the same circuit.")
    add_body(doc, "The post-qualifying information set added grid position, qualifying position, and gap to pole. These variables are operationally available only after qualifying and were never supplied to pre-qualifying models. Feature context was treated as an explicit contract: a model trained in one context was not allowed to receive columns available only in the other context.")
    add_body(doc, "Weather fields were excluded from the thesis candidates. The design review found that the then-current target-race weather aggregation was not bounded by a pre-race timestamp or forecast provenance and could therefore leak race-time information. Imputation, scaling, class weighting, and threshold selection were fit within the relevant training partitions. This prevents information from a validation or held-out season from influencing preprocessing choices.")
    add_body(doc, "Two remaining implementation risks should be retained in the validity discussion. First, pre-qualifying entry and team mapping uses the latest prior completed race, so late substitutions or team changes may be missed. Second, historical post-qualifying rows use recorded grid position, whereas upcoming inference may use qualifying position as a proxy unless an official grid is available. These are parity risks rather than evidence of predictive effect.")

    add_heading(doc, "4. Four Prediction Tasks and Their Metrics", 1)
    add_table(doc, ["Task", "Target definition", "Primary metric", "Key secondary evidence"], [
        ["Finishing position", "Classified finishing position", "MAE", "RMSE, R2, mean per-race Spearman rank correlation."],
        ["Top 10", "1 if finishing position is at most 10", "ROC-AUC", "PR-AUC, F1, precision, recall, balanced accuracy, Brier score, log loss."],
        ["Podium", "1 if finishing position is at most 3", "PR-AUC", "ROC-AUC, F1, precision, recall, balanced accuracy, Brier score, log loss."],
        ["Position gain/loss", "Grid position minus finishing position", "MAE", "RMSE, R2, and sign accuracy."],
    ], [1740, 2740, 1440, 3440])
    add_caption(doc, "Suggested Table 2. Prediction tasks and evaluation criteria. Lower values are better for MAE, RMSE, Brier score, and log loss; higher values are better for the remaining metrics.")
    add_body(doc, "MAE was selected for the regression tasks because it expresses average error directly in positions. RMSE exposes larger errors, R2 describes explained variation, and mean per-race Spearman correlation assesses finishing-order quality within races. For position gain/loss, sign accuracy is reported because an error can be small in magnitude while still predicting the wrong direction.")
    add_body(doc, "For Top 10 classification, ROC-AUC was used as the primary discrimination measure. For podium classification, PR-AUC was primary because podium finishes form the minority class and raw accuracy can be misleading. Brier score and log loss were retained as probability-quality measures. Classification thresholds were selected on an earlier inner validation season and applied unchanged to outer validation and the 2025 holdout.")

    add_heading(doc, "5. Candidate Algorithms, Baselines, and Tuning Protocol", 1)
    add_body(doc, "Each context-task combination was modelled separately. Regression comparisons included a constant median baseline, Ridge, ElasticNet, Random Forest Regressor, XGBoost Regressor, and LightGBM Regressor. The post-qualifying finishing-position task also included a grid-position operational baseline. Classification comparisons included a prevalence baseline, Logistic Regression, Random Forest Classifier, XGBoost Classifier, and LightGBM Classifier. Position gain/loss used a zero-change baseline, which is substantively meaningful because no change is a plausible default forecast.")
    add_body(doc, "Candidate configurations were seeded and received the same temporal split and approved feature set within each context-task comparison. Linear models used standardized numerical features; tree models received the same fold-local imputed inputs. The validation winner was selected by the task's primary metric, with a simpler-model tie rule. The final 2025 season was not used to replace a validation-selected winner after the fact.")
    add_body(doc, "The actual experiment is appropriately interpreted as a fixed, reproducible candidate comparison rather than proof that its specific hyperparameter settings are globally optimal. The single validation fold makes any fine-grained ranking among close candidates especially uncertain.")

    add_heading(doc, "6. Experimental Results and Interpretation", 1)
    add_table(doc, ["Context", "Task", "Validation champion", "Primary validation score", "2025 held-out result"], [
        ["Pre", "Finishing position", "Ridge", "MAE 3.426", "MAE 3.904; R2 0.286; rank correlation 0.510."],
        ["Post", "Finishing position", "Ridge", "MAE 2.804", "MAE 3.258; R2 0.452; rank correlation 0.649."],
        ["Pre", "Top 10", "Logistic Regression", "ROC-AUC 0.853", "ROC-AUC 0.759; F1 0.679; Brier 0.199."],
        ["Post", "Top 10", "Logistic Regression", "ROC-AUC 0.912", "ROC-AUC 0.837; F1 0.766; Brier 0.164."],
        ["Pre", "Podium", "Random Forest", "PR-AUC 0.494", "PR-AUC 0.523; F1 0.578; Brier 0.097."],
        ["Post", "Podium", "Random Forest", "PR-AUC 0.689", "PR-AUC 0.750; F1 0.734; Brier 0.067."],
        ["Pre", "Gain/loss", "Zero-change baseline", "MAE 2.898", "MAE 3.322; R2 approximately 0; sign accuracy 0.182."],
        ["Post", "Gain/loss", "ElasticNet", "MAE 2.797", "MAE 3.261; R2 0.219; sign accuracy 0.544."],
    ], [720, 1420, 1700, 1660, 3860])
    add_caption(doc, "Suggested Table 3. Validation-selected champions and final 2025 held-out results. Values are reported by context and task; the holdout was not used to select the champion.")
    add_body(doc, "The results do not indicate a universal advantage for complex learners. Ridge was selected for finishing position in both contexts, and Logistic Regression was selected for Top 10 classification in both contexts. Random Forest was selected for podium classification. This pattern is useful because it demonstrates that the contribution is the fair temporal comparison rather than an assumption that boosted or ensemble models must dominate.")
    add_body(doc, "Position gain/loss is the most limited task. The pre-qualifying zero-change baseline was selected on validation and had held-out R2 approximately zero with sign accuracy 0.182. The post-qualifying ElasticNet was stronger, but its MAE improvement over pre-qualifying was small. These outcomes should be presented as a negative or weak result rather than reframed as broad position-change predictability.")
    add_caption(doc, "Suggested Figure 1. Model leaderboards by task, generated from the saved aggregate results: `figures/leaderboard_position_model.png`, `leaderboard_top10_model.png`, `leaderboard_podium_model.png`, and `leaderboard_position_gain_model.png`.")

    add_heading(doc, "7. Pre- versus Post-Qualifying Analysis", 1)
    add_body(doc, "The 2025 held-out comparison indicates a consistent advantage for the post-qualifying context in finishing position, Top 10 classification, and podium classification. Finishing-position MAE declined from 3.904 to 3.258 positions, while mean per-race rank correlation increased from 0.510 to 0.649. Top 10 ROC-AUC increased from 0.759 to 0.837 and the Brier score declined from 0.199 to 0.164. The largest difference occurred for podium PR-AUC, which increased from 0.523 to 0.750; the corresponding Brier score declined from 0.097 to 0.067.")
    add_body(doc, "These results are consistent with the additional predictive information contained in grid position, qualifying position, and gap to pole. They do not establish a causal effect of qualifying on race outcomes, nor do they imply that the same magnitude of improvement will occur in all future seasons. Position gain/loss provides the necessary counterexample: its MAE changed only from 3.322 to 3.261 positions, although post-qualifying sign accuracy improved to 0.544.")
    add_caption(doc, "Suggested Figure 2. Context comparison of validation champions: `figures/context_performance_comparison.png`. The caption should identify the metric for each task because the score scales differ across tasks.")

    add_heading(doc, "8. Feature-Ablation Analysis", 1)
    add_body(doc, "Feature ablations compared driver/team form only, form plus circuit history, and the full context-specific feature set. The post-qualifying full set, including grid and qualifying information, was best for every task in validation. For finishing position, its MAE was 2.804 compared with 3.422 for form only. For Top 10, ROC-AUC was 0.912 compared with 0.848 for form only. For podium, PR-AUC was 0.688 compared with 0.469. For gain/loss, MAE improved modestly from 2.906 for the form-only zero-change winner to 2.797 for the full ElasticNet specification.")
    add_body(doc, "The pre-qualifying picture was more restrained. Form only was marginally best for finishing position (MAE 3.423 versus 3.426 for all features) and Top 10 (ROC-AUC 0.853 in both practical terms). Full pre-qualifying features were best for podium PR-AUC (0.507 versus 0.473 for form only). For gain/loss, every ablation retained the zero-change baseline with MAE 2.898. Thus, the ablations support an association between the additional post-qualifying information and performance, while showing that not every historical feature improves every task.")
    add_caption(doc, "Suggested Figure 3. Feature-ablation comparison generated from `figures/feature_ablation_comparison.png`. Interpret each bar as a validation-selected model within a feature subset, not as a causal contribution estimate.")

    add_heading(doc, "9. Limitations, Validity Threats, and Future Work", 1)
    add_body(doc, "Several limitations constrain the interpretation of these findings. Formula 1 outcomes are affected by safety cars, red flags, collisions, reliability failures, strategy changes, and weather evolution during a race. These factors are either unobserved or uncertain at the chosen prediction cutoff. Consequently, the models should be treated as analytical forecasts rather than causal explanations, betting advice, or guarantees of performance.")
    add_body(doc, "The most immediate internal-validity limitation is the single outer validation fold caused by absent 2022 race-result data. The 2025 holdout is genuinely later and untouched, but a one-fold development comparison gives limited evidence about season-to-season stability. Podium events are also imbalanced, which is why PR-AUC and probability-quality measures were emphasised; however, performance estimates may still be sensitive to a small number of positive cases.")
    add_body(doc, "Feature provenance and deployment parity require continued attention. Target-race weather fields were excluded because their pre-race availability could not be established. Historical grid position and live qualifying-position proxy behaviour must remain aligned. Missing lap data, driver substitutions, team changes, and circuit-name matching can introduce measurement error or selection bias. Competitive order and technical regulations can also change across seasons, limiting generalisation beyond the documented data snapshot.")
    add_bullets(doc, [
        "Future work should restore complete 2022 result coverage and add further completed seasons to obtain multiple outer folds and more stable uncertainty estimates.",
        "A timestamped weather forecast source should be added before weather variables are reconsidered for thesis modelling.",
        "Race-cluster bootstrap intervals and calibration analyses should be expanded in later replications, particularly for the rare podium class.",
        "Live inference should use an official starting grid when available, or the training and inference contracts should consistently use the same qualifying-position proxy.",
    ])

    add_heading(doc, "Abstract-Style Contribution Summary", 1)
    add_body(doc, "This study presents a reproducible, leakage-aware comparison of machine-learning models for four driver-level Formula 1 prediction tasks: finishing position, Top 10 finish, podium finish, and position gain or loss. Models were evaluated in pre-qualifying and post-qualifying information contexts using completed seasons only. The final design trained on 2021, 2023, and 2024, selected models through a chronological validation fold ending in 2024, and evaluated frozen champions on the held-out 2025 season. Post-qualifying models outperformed pre-qualifying models for finishing position, Top 10 classification, and podium classification: held-out finishing-position MAE decreased from 3.904 to 3.258 positions, Top 10 ROC-AUC increased from 0.759 to 0.837, and podium PR-AUC increased from 0.523 to 0.750. Ridge, Logistic Regression, and Random Forest won different tasks, indicating that model complexity was not uniformly beneficial. Position gain/loss remained difficult before qualifying, where a zero-change baseline was selected. The findings support context-sensitive analytical forecasting, while acknowledging limited validation folds, race incidents, strategy, weather uncertainty, class imbalance, and season-to-season regulatory change.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
