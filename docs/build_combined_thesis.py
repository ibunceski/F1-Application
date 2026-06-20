"""Build the combined, ML-centred ApexInsight thesis draft.

Sources are deliberately local and traceable: the original platform report,
the thesis ML chapter, and the final experiment/design markdown reports.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
import shutil

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
SOURCE = DOCS / "ApexInsight - F1 Prediction Platform.docx"
OUT = DOCS / "ApexInsight-ML-Thesis-Draft.docx"
ASSETS = DOCS / ".thesis_assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
PALE = "F4F6F9"
TABLE_WIDTH = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=11, color=INK, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("ApexInsight | F1 ML Thesis")
    set_run_font(run, size=9, color=MUTED, italic=True)
    add_page_number(section.footer.paragraphs[0])


def add_body(doc, text: str, *, centered=False, italic=False, small=False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_run_font(r, size=9.5 if small else 11, color=INK if not small else MUTED, italic=italic)


def add_bullets(doc, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_numbered(doc, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def add_table(doc, headers, rows, widths, *, font_size=9.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(label))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, values)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(values)-1) and len(str(value)) < 20 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)


def extract_source_images() -> list[Path]:
    if ASSETS.exists():
        shutil.rmtree(ASSETS)
    ASSETS.mkdir(parents=True)
    with ZipFile(SOURCE) as archive:
        names = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
        for name in names:
            output = ASSETS / Path(name).name
            output.write_bytes(archive.read(name))
    return sorted(ASSETS.iterdir(), key=lambda path: path.name)


def add_source_figure(doc, images, index, caption, width=5.85) -> None:
    if index < len(images):
        image_path = images[index]
        with Image.open(image_path) as source_image:
            pixel_width, pixel_height = source_image.size
        computed_height = width * pixel_height / pixel_width
        max_height = 7.15
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        if computed_height > max_height:
            p.add_run().add_picture(str(image_path), height=Inches(max_height))
        else:
            p.add_run().add_picture(str(image_path), width=Inches(width))
        add_caption(doc, caption)
    else:
        add_caption(doc, caption + " (илустрацијата е достапна во изворната проектна документација.)")


def heading(doc, text: str, level=1, page_break=False) -> None:
    if page_break:
        doc.add_page_break()
    doc.add_heading(text, level=level)


def build() -> None:
    images = extract_source_images()
    doc = Document()
    configure_document(doc)

    # Cover: editorial_cover form factor, adapted to the academic thesis.
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("ДИПЛОМСКА РАБОТА — РАБОТЕН НАЦРТ")
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ApexInsight")
    set_run_font(r, size=30, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Платформа за анализа на Formula 1 и репродуцибилна споредба на модели за предвидување")
    set_run_font(r, size=15, color=DARK_BLUE)
    add_body(doc, "Машинското учење е истражувачкото јадро; веб-платформата обезбедува интеграција, проверливост и аналитичка презентација.", centered=True, italic=True)
    for _ in range(6):
        doc.add_paragraph()
    add_body(doc, "Факултет за информатички науки и компјутерско инженерство", centered=True, small=True)
    add_body(doc, "Скопје, јуни 2026", centered=True, small=True)

    heading(doc, "Апстракт", page_break=True)
    add_body(doc, "Оваа работа го претставува ApexInsight, платформа за анализа на Formula 1 податоци чиј главен научен придонес е репродуцибилна и временски коректна споредба на модели за четири предвидувачки задачи на ниво возач–трка: завршна позиција, пласман во првите десет, пласман на подиум и добивка/загуба на позиции. Платформата обединува прибирање податоци, PostgreSQL модел, feature engineering, FastAPI сервис и React интерфејс, но експерименталниот дизајн е поставен така што заклучоците не зависат од веб-демонстрацијата. Моделите се споредени во pre-qualifying и post-qualifying контекст, со строго хронолошка валидација и со 2025 како недопрена завршна сезона. Не се користат нецелосни податоци од 2026, а временските полиња на целната трка се исклучени поради недоволно докажано pre-race потекло. На final holdout сезоната post-qualifying моделите постигнаа MAE 3.258 за завршна позиција, ROC-AUC 0.837 за Top 10 и PR-AUC 0.750 за podium; pre-qualifying gain/loss задачата не покажа убедлива научена предност над zero-change базната линија. Резултатите ја поддржуваат употребата на контекстно зависни аналитички прогнози, но не претставуваат каузални објаснувања или гаранции за идни трки.")
    add_body(doc, "Клучни зборови: Formula 1, машинско учење, временска валидација, предвидување, FastAPI, React, PostgreSQL, репродуцибилност.", italic=True)

    heading(doc, "Содржина", page_break=True)
    toc = [
        "1. Вовед", "2. Истражувачка цел и опфат", "3. Архитектура на системот", "4. Технологии и алатки",
        "5. Модел на податоци", "6. Прибирање и обработка на податоци", "7. Машинско учење и предвидувања",
        "8. Backend API", "9. Frontend апликација", "10. Главни функционалности", "11. Генерирање и користење предвидувања",
        "12. Репродуцибилност и верификација", "13. Model Lab: презентација на експериментални докази",
        "14. Ограничувања", "15. Идни подобрувања", "16. Заклучок",
    ]
    add_numbered(doc, toc)
    add_body(doc, "Напомена: насловите се дефинирани со Word heading стилови; автоматската содржина со броеви на страници може да се освежи во Microsoft Word по финалното уредување.", small=True, italic=True)

    heading(doc, "1. Вовед", page_break=True)
    add_body(doc, "Formula 1 претставува погоден домен за анализирање на податоци затоа што секоја трка создава истовремено спортски, технички и временски зависни набљудувања. Завршниот поредок на возачот е поврзан со историската форма, перформансите на тимот, карактеристиките на патеката, квалификацискиот резултат и стартната позиција, но и со настани што не се познати пред трката. Таквата комбинација го прави доменот корисен за машинско учење, но и методолошки чувствителен: случајна поделба на редовите или користење информации по трката може да создаде привидно висока, но неупотреблива точност.")
    add_body(doc, "ApexInsight е full-stack платформа што ги поврзува историските Formula 1 податоци со предвидувачки модели и аналитички интерфејс. Претходната проектна верзија беше ориентирана првенствено кон функционалностите на веб-апликацијата. Во оваа работа фокусот се преместува кон систематска споредба на кандидати за модели: истите задачи, истите target дефиниции, истите временски поделби и еднакви feature контракти се применуваат на сите алгоритми. Платформата останува важна затоа што ги прави податоците, моделите и артефактите достапни и проверливи, но не е замена за експериментален доказ.")
    add_body(doc, "Работата не е betting систем и не тврди дека може да го предвиди целосниот тек на една трка. Нејзината цел е да покаже како инженерски систем може да поддржи строг ML експеримент: од вчитување и чистење податоци, преку feature engineering со временски ограничувања, до зачувани резултати, API и визуелна анализа.")

    heading(doc, "2. Истражувачка цел и опфат")
    add_body(doc, "Истражувачката цел е да се утврди дали конвенционални линеарни и tree-based модели даваат корисни driver-level Formula 1 предвидувања кога им се дозволени само информации достапни пред трката. Главниот придонес е репродуцибилната, leakage-safe споредба на повеќе кандидатски семејства во два оперативни информациски контексти, а не тврдење дека еден алгоритам е универзално најдобар.")
    add_bullets(doc, [
        "RQ1: Кои кандидатски семејства ги подобруваат значајните базни линии во четирите задачи при хронолошка валидација?",
        "RQ2: Во која мера post-qualifying информацискиот сет ја подобрува изведбата во однос на pre-qualifying сетот?",
        "RQ3: Дали моделите избрани на валидациските сезони задржуваат корисна изведба и калибрација на една подоцнежна, целосно завршена сезона?",
        "RQ4: Дали веројатностите за Top 10 и podium имаат доволна дискриминација и калибрација за аналитичкото прикажување во апликацијата?",
    ])
    add_body(doc, "Единицата на анализа е ред за возач–трка. Бидејќи повеќе возачи учествуваат во иста трка, тие не се третираат како независни набљудувања во интерпретацијата на неизвесноста; резултатите по трка се задржуваат во артефактите за понатамошна кластерска анализа.")

    heading(doc, "3. Архитектура на системот", page_break=True)
    add_body(doc, "Системот е изграден како full-stack архитектура со одвоени одговорности. Историските Formula 1 податоци се внесуваат преку ingestion процеси, се складираат во PostgreSQL, се трансформираат во feature редови, а потоа се користат од експерименталниот runner и prediction сервисот. FastAPI го изложува доменскиот API и read-only Model Lab API, а React клиентот ги презентира и оперативните предвидувања и артефактите од споредбата на модели.")
    add_bullets(doc, [
        "db: PostgreSQL 16, со нормализирани ентитети и ML feature/prediction записи;",
        "backend: FastAPI апликација со REST endpoints, валидација и model loading;",
        "frontend: React/Vite клиент со TypeScript, React Query и визуелни компоненти;",
        "ingestion: процес што се активира по потреба за преземање, обработка и feature engineering;",
        "experiment runner: репродуцибилен процес што пишува immutable artifacts и промовира модели само по успешно завршување.",
    ])
    add_source_figure(doc, images, 0, "Слика 1. Архитектурен тек: од извори на податоци до API, модели и dashboard. Адаптирано од проектната документација.")
    add_body(doc, "Од аспект на оваа теза, најважната граница е меѓу историското складирање и одлуките донесени на prediction cutoff. Feature engineering не смее да чита податоци од целната трка ако тие не биле достапни во соодветниот контекст. Исто така, експерименталните артефакти се одвоени од deployed joblib датотеките, за неуспешен или нецелосен run да не ја промени апликацијата.")

    heading(doc, "4. Технологии и алатки")
    add_body(doc, "Backend делот е имплементиран со FastAPI и Python. FastAPI обезбедува Pydantic валидација, OpenAPI документација и јасни dependency граници; SQLAlchemy 2.0 и Alembic се користат за пристап и миграции на базата. Централизираното справување со грешки и X-Request-ID ја олеснуваат дијагностиката на API повиците.")
    add_body(doc, "За податоци и машинско учење се користат FastF1, pandas, NumPy, scikit-learn, XGBoost, LightGBM и joblib. Моделите се поставени како fold-local pipeline-и со имputation и, кога е потребно, scaling. Ова е важно не само поради квалитет на моделот, туку и поради спречување истекување од валидациски или holdout редови кон трансформациите научени при тренинг.")
    add_body(doc, "Frontend-от е изграден со React, TypeScript и Vite; TanStack React Query се користи за API cache и состојби на вчитување, Recharts за интерактивни графикони, а Tailwind CSS за конзистентна визуелна хиерархија. Docker Compose ги координира сервисите и овозможува повторливо локално стартување.")

    heading(doc, "5. Модел на податоци", page_break=True)
    add_body(doc, "Релацискиот модел е поставен околу сезоните, трките, возачите, тимовите и резултатите. Табелите seasons и races обезбедуваат временска и календарска структура; drivers и teams го чуваат идентитетот на учесниците; race_results и qualifying_results ги содржат исходите и квалификациските записи. lap_times и weather_data се сурови детални извори, додека ml_features и predictions се наменети за ML workflow-от и апликациските предвидувања.")
    add_table(doc, ["Табела", "Улога во тезата", "Клучна контрола"], [
        ["seasons / races", "Календар, редослед и дати на настаните", "Проверка дека сезоната е целосна пред евалуација"],
        ["race_results", "Фиксни target вредности и grid позиција", "Нема импутирање на недостасни target labels"],
        ["qualifying_results", "Квалификациски позиција и gap до pole", "Достапни само во post-qualifying контекст"],
        ["lap_times", "Историски clean-lap темпо", "Се користат само претходни трки"],
        ["ml_features", "Feature редови по возач–трка и контекст", "Контекстот е изречен feature contract"],
        ["predictions", "Оперативни предвидувања и верзија на модел", "Контекст и моделска верзија се зачувани"],
    ], [1800, 4700, 2860])
    add_source_figure(doc, images, 1, "Слика 2. Концептуален приказ на релациите во базата на податоци. Адаптирано од проектната документација.")

    heading(doc, "6. Прибирање и обработка на податоци", page_break=True)
    add_body(doc, "Ingestion скриптите ги внесуваат основните информации за сезони, трки, возачи и тимови, потоа резултатите и квалификациите, а по потреба и круговите, гумите и временските записи. Овие чекори не се поистоветуваат со ML тренинг: тие создаваат versionable data snapshot од кој потоа се градат feature редови и експериментални поделби.")
    add_body(doc, "Feature engineering користи претходни трки за историското темпо, recent form, team form, circuit history и DNF стапки. За pre-qualifying контекстот не се додаваат grid, qualifying position или gap to pole. Post-qualifying контекстот ги дополнува истите историски сигнали со овие три променливи, бидејќи се достапни откако ќе завршат квалификациите и е познат предтркачкиот grid.")
    add_body(doc, "Во тезинските експерименти weather_is_wet и avg_track_temp_c се исклучени. Во постојната агрегација нивното потекло за целната трка не е ограничено со pre-race timestamp или докажана forecast provenance; нивно вклучување би можело да внесе race-time информација. Ова е намерна feature-policy корекција, а не промена на целните променливи.")
    add_bullets(doc, [
        "2021, 2023, 2024 и 2025 се прифатени како сезони со комплетна race и qualifying покриеност;",
        "2022 е исклучена бидејќи има qualifying записи, но нема race-result редови;",
        "2026 е исклучена од сите тезински поделби бидејќи содржи идни закажани трки и делумни резултати;",
        "секоја сезона за евалуација мора да мине fail-closed проверка за број на трки, резултати, квалификации и немање идни настани.",
    ])

    heading(doc, "7. Машинско учење и предвидувања", page_break=True)
    heading(doc, "7.1 Фиксни задачи, targets и метрики", level=2)
    add_body(doc, "Четирите targets се задржани непроменети од апликацијата. Завршна позиција е ingested classified finishing_position. Top 10 е 1 ако finishing_position е најмногу 10, podium е 1 ако е најмногу 3, а position gain/loss е grid_position минус finishing_position; позитивна вредност значи добиени позиции. Редовите без неопходен target не се пополнуваат вештачки и не се вклучуваат во соодветната задача.")
    add_table(doc, ["Задача", "Primary metric", "Secondary evidence"], [
        ["Завршна позиција", "MAE", "RMSE, R², mean per-race Spearman"],
        ["Top 10", "ROC-AUC", "PR-AUC, F1, precision, recall, balanced accuracy, Brier, log loss"],
        ["Podium", "PR-AUC", "ROC-AUC, F1, precision, recall, balanced accuracy, Brier, log loss"],
        ["Gain/loss", "MAE", "RMSE, R², sign accuracy"],
    ], [2450, 2050, 4860])
    add_body(doc, "MAE се користи како примарна регресиска метрика бидејќи грешката се чита директно во позиции. Mean per-race Spearman ја оценува способноста за рангирање во рамки на една трка. За podium, PR-AUC е примарна поради ретката позитивна класа; за класификациите Brier score и log loss ја дополнуваат дискриминацијата со квалитет на веројатностите. Праговите за F1, precision и recall се избираат само од валидациски предвидувања и остануваат замрзнати на holdout сезоната.")

    heading(doc, "7.2 Кандидатски модели, базни линии и tuning protocol", level=2)
    add_body(doc, "Секој пар context–task добива посебна споредба со ист временски split и ист одобрен feature set. За регресија се споредуваат median baseline, Ridge, ElasticNet, Random Forest, XGBoost и LightGBM. За finishing position во post-qualifying контекст дополнително се известува grid-position operational baseline. За класификација се споредуваат prevalence/dummy baseline, Logistic Regression, Random Forest, XGBoost и LightGBM. За gain/loss zero-change baseline е суштинска референтна точка, затоа што ненадејната промена на позиција не е подразбрана.")
    add_body(doc, "Сите fixed configurations се seeded и ја користат истата fold-local preprocessing процедура. Победникот се избира според примарната метрика на validation fold; во практично изедначување важи правило за поедноставен модел. Финалната 2025 сезона не се користи за избор на алгоритам, feature subset, hyperparameter или threshold. Така, final evaluation е потврдна, а не селекциска фаза.")
    add_source_figure(doc, images, 2, "Слика 3. Data-flow на training pipeline. Сликата го прикажува инженерскиот тек; тезинската селекција користи само завршени сезони и immutable experiment artifacts.")
    add_source_figure(doc, images, 3, "Слика 4. Метрики за оценување. Во конечниот дизајн MAE е primary за регресија, ROC-AUC за Top 10 и PR-AUC за podium.")

    heading(doc, "7.3 Временска поделба и спречување leakage", level=2)
    add_body(doc, "Валидацијата е season-based rolling origin. Бидејќи недостигаат резултати за 2022, достапен е еден најран изводлив outer fold: тренинг на 2021 и 2023, валидација на 2024. Потоа избраните конфигурации се refit-ираат на 2021, 2023 и 2024 и се оценуваат еднаш на 2025. Ова е reduced-evidence дизајн, но е построг од random cross-validation и не дозволува идни трки или сезони да влијаат врз претходен fold.")
    add_table(doc, ["Фаза", "Сезони", "Улога"], [
        ["Развој", "2021, 2023, 2024", "Достапни целосни сезони пред holdout"],
        ["Rolling validation", "train: 2021, 2023 → validate: 2024", "Избор на model family, threshold и feature subset"],
        ["Final holdout", "2025", "Една недопрена потврдна евалуација"],
    ], [1800, 3000, 4560])
    add_body(doc, "Feature пресметките се ограничени на трки порано од целната трка. Imputation, scaling, class weighting и threshold избор се fit-ираат само на training делот од релевантниот fold. Pre-qualifying модел никогаш не прима колони што се достапни само после квалификации. Овие контроли се проверуваат пред креирање на split и се запишуваат во manifest за секој experiment ID.")

    heading(doc, "7.4 Конечен експеримент и резултати", level=2)
    add_body(doc, "Финалниот успешен run е thesis-final-2025-holdout-20260620-r3, со seed 42. Во snapshot-от имало 420/440/479/479 pre-qualifying rows и 439/440/479/479 post-qualifying rows за 2021/2023/2024/2025. Не се импутираат недостасни labels. По успешното создавање на артефактите се промовирани само validation-selected champion модели со постојните deployed filenames, а постарите experiment директориуми не се презапишани.")
    add_table(doc, ["Контекст", "Задача", "Champion", "Validation metric", "2025 holdout"], [
        ["Pre", "Завршна позиција", "Ridge", "MAE 3.426", "MAE 3.904; R² 0.286; ρ 0.510"],
        ["Post", "Завршна позиција", "Ridge", "MAE 2.804", "MAE 3.258; R² 0.452; ρ 0.649"],
        ["Pre", "Top 10", "Logistic Regression", "ROC-AUC 0.853", "ROC-AUC 0.759; F1 0.679; Brier 0.199"],
        ["Post", "Top 10", "Logistic Regression", "ROC-AUC 0.912", "ROC-AUC 0.837; F1 0.766; Brier 0.164"],
        ["Pre", "Podium", "Random Forest", "PR-AUC 0.494", "PR-AUC 0.523; F1 0.578; Brier 0.097"],
        ["Post", "Podium", "Random Forest", "PR-AUC 0.689", "PR-AUC 0.750; F1 0.734; Brier 0.067"],
        ["Pre", "Gain/loss", "Zero-change baseline", "MAE 2.898", "MAE 3.322; R² ≈ 0; sign 0.182"],
        ["Post", "Gain/loss", "ElasticNet", "MAE 2.797", "MAE 3.261; R² 0.219; sign 0.544"],
    ], [800, 1600, 2100, 1900, 2960], font_size=8.2)
    add_body(doc, "Резултатите не покажуваат универзална предност на комплексните алгоритми: Ridge победува за завршна позиција, Logistic Regression за Top 10, а Random Forest за podium. Најважниот негативен резултат е pre-qualifying gain/loss, каде zero-change baseline е избран на валидација и holdout R² е приближно нула. Оваа задача треба да се интерпретира како ограничена, а не како доказ за сигурно предвидување на промени во позиции.")
    add_source_figure(doc, images, 4, "Слика 5. Пример за feature-importance приказ во платформата. За тезата, feature importance е објаснувачка помош, а не критериум за избор на champion.")

    heading(doc, "7.5 Pre-qualifying наспроти post-qualifying", level=2)
    add_body(doc, "На holdout сезоната post-qualifying контекстот е подобар во три главни задачи. MAE за завршна позиција се намалува за 0.646 позиции, од 3.904 на 3.258, а mean per-race rank correlation расте од 0.510 на 0.649. За Top 10 ROC-AUC се зголемува од 0.759 на 0.837, со пад на Brier score од 0.199 на 0.164. Најголемата разлика е кај podium: PR-AUC расте од 0.523 на 0.750, а Brier score се намалува од 0.097 на 0.067.")
    add_body(doc, "Овие разлики се конзистентни со дополнителната информативност на grid position, qualifying position и gap to pole. Тие не докажуваат каузален ефект на квалификациите врз исходот и не гарантираат иста добивка во друга сезона. Gain/loss е контрапример: MAE се менува само од 3.322 на 3.261, иако post-qualifying sign accuracy расте на 0.544.")

    heading(doc, "7.6 Feature-ablation анализа", level=2)
    add_body(doc, "Ablation експериментите ги споредуваат driver/team form само, form плус circuit history и целосниот сет за соодветниот контекст. Секоја ablation споредба го користи истиот временски fold и кандидатска матрица, па претставува опис на разликите во predictive performance, а не каузална проценка на поединечните features.")
    add_table(doc, ["Контекст", "Задача", "Најдобар subset", "Metric / score", "Клучна споредба"], [
        ["Post", "Завршна позиција", "All + grid/qualifying", "MAE 2.804", "form only: 3.422"],
        ["Post", "Top 10", "All + grid/qualifying", "ROC-AUC 0.912", "form only: 0.848"],
        ["Post", "Podium", "All + grid/qualifying", "PR-AUC 0.688", "form only: 0.469"],
        ["Post", "Gain/loss", "All + grid/qualifying", "MAE 2.797", "zero/form only: 2.906"],
        ["Pre", "Завршна позиција", "Form only", "MAE 3.423", "all: 3.426"],
        ["Pre", "Top 10", "Form only", "ROC-AUC 0.853", "all: 0.853"],
        ["Pre", "Podium", "All", "PR-AUC 0.507", "form only: 0.473"],
        ["Pre", "Gain/loss", "Any; zero baseline", "MAE 2.898", "нема демонстрирана корист"],
    ], [800, 1500, 2250, 1800, 3010], font_size=8.2)
    add_body(doc, "Полниот post-qualifying сет е најдобар за сите задачи на валидација. Pre-qualifying сликата е повнимателна: form-only е практично изедначен или подобар за завршна позиција и Top 10, додека сите pre-qualifying features имаат умерена предност само за podium. Ова е во согласност со контекстно зависна корист од дополнителните информации, а не со претпоставка дека повеќе features секогаш помагаат.")

    heading(doc, "8. Backend API", page_break=True)
    add_body(doc, "Backend API-то е имплементирано со FastAPI под /api/v1. Домeнските групи /seasons, /races, /drivers, /teams и /analysis обезбедуваат пристап до историските и аналитичките податоци. Групата /predictions ги задржува постојните операции за генерирање и преглед на предвидувања, model-info и feature importance, со цел да остане компатибилна со постојниот frontend.")
    add_body(doc, "За тезинскиот workflow е додаден read-only /model-lab API. Тој не тренира модели; чита само persistent artifacts од models_store/experiments/<experiment_id>. Endpoints овозможуваат листа на успешни експерименти, overview со методологија и champion-и, филтрирани candidate results, ablation резултати и metadata за достапните figures. Ако не е наведено experiment ID, сервисот безбедно избира latest successful експеримент; нецелосен или malformed artifact дава јасен validation/404 одговор.")
    add_bullets(doc, [
        "GET /model-lab/experiments — идентификатор, timestamp, контексти, evaluation сезони и статус;",
        "GET /model-lab/overview — конфигурација, data summary, champions и aggregate leaderboard;",
        "GET /model-lab/results — филтрирање по задача, контекст, алгоритам и evaluation сезона;",
        "GET /model-lab/ablations и /model-lab/artifacts — споредби на feature subset-и и metadata за CSV/figure артефакти.",
    ])
    add_body(doc, "Централизираното error handling враќа структурирана грешка, message, request_id и timestamp. Во production може да се користи X-API-Key, а health endpoint-от ја известува достапноста на сервисот, базата и вчитаните модели.")

    heading(doc, "9. Frontend апликација", page_break=True)
    add_body(doc, "Frontend-от е React апликација со TypeScript. AppShell обезбедува sidebar, top bar и заедничка структура; страници се организирани според сезонски преглед, конкретна трка, анализа, гуми, споредба на возачи и предвидувања. React Query управува со cache, loading и error состојби, додека типовите од API слојот го ограничуваат несогласувањето меѓу backend и интерфејс.")
    add_body(doc, "Model Lab е поставен како посебна, истакната рута (/model-lab) и не смее да се меша со страницата „Prediction versus actual outcome“. Првата одговара на истражувачкото прашање „кој модел е избран под каков период и контекст?“, додека втората е оперативен преглед на предвидување за една конкретна трка. Кога недостигаат артефакти, интерфејсот прикажува informative empty/error state, без да имплицира дека нема модел или дека неуспешен run е доказ.")
    add_bullets(doc, [
        "Experiment selector и context filter за избор на запишан, успешен run;",
        "task tabs за четирите цели, sortable leaderboard и јасно означен champion;",
        "tooltip-и за MAE, PR-AUC, ROC-AUC, Brier score и rank correlation;",
        "контекстна споредба, ablation приказ и линк од победнички модел кон feature-importance објаснувањето;",
        "при секое тврдење за изведба се прикажуваат metric, evaluation period и feature context.",
    ])

    heading(doc, "10. Главни функционалности", page_break=True)
    heading(doc, "10.1 Сезонски и тркачки аналитички погледи", level=2)
    add_body(doc, "Season overview прикажува календар, standings и основни статистики. Race selector го поврзува сезонскиот приказ со деталните страници за конкретен настан. Race Analysis обединува резултати, најбрзи кругови, временски услови, промени на позиции и lap-time графикони. Tyre Strategy ги прикажува compound-ите, stint-овите и pit stop-овите, а Driver Comparison овозможува head-to-head анализа на двајца возачи.")
    add_source_figure(doc, images, 5, "Слика 6. Пример за Race Analysis интерфејс. Адаптирано од проектната документација.")
    add_source_figure(doc, images, 6, "Слика 7. Пример за Tyre Strategy интерфејс. Адаптирано од проектната документација.")
    heading(doc, "10.2 Предвидувачки и истражувачки погледи", level=2)
    add_body(doc, "Race Predictor генерира и прикажува предвидена позиција, Top 10 и podium веројатности, очекувана промена на позиција и confidence score за избраната трка и контекст. Next Race Prediction е насочен кон најблискиот настан. Prediction Comparison е достапен кога постојат реални резултати и визуелно ги споредува оперативните предвидувања со исходот. Model Lab ја дополнува оваа функционалност со ниво на анализа што не е врзано за една трка: кандидатски leaderboard-и, champion-и, validation/holdout периоди и feature ablation-и.")

    heading(doc, "11. Генерирање и користење предвидувања", page_break=True)
    add_body(doc, "Кога корисникот бара предвидување, backend сервисот проверува дали постојат потребните feature rows за избраната трка и контекст, го избира вчитаниот champion модел и пишува резултат во predictions. Записот носи race, driver, prediction context, model/feature context, верзија на модел, предвидена позиција, веројатности и време на генерирање. На тој начин не се мешаат предвидувања направени пред и по квалификации.")
    add_body(doc, "Промоцијата на нови модели е одделена од самото генерирање предвидувања. Experiment runner прво креира immutable manifest, CSV резултати, row-level out-of-fold predictions, калибрациски податоци и Markdown report. Само ако run-от заврши успешно, осумте избрани champion-и (четири задачи за два контекста) се зачувуваат со постојните deployed joblib имиња. Со тоа prediction API-то останува backward-compatible, а Model Lab може да објасни од кој експеримент потекнува активниот модел.")
    add_source_figure(doc, images, 7, "Слика 8. Пример за Race Predictor интерфејс. Веројатностите треба да се читаат заедно со контекстот, калибрациските метрики и evaluation периодот.")

    heading(doc, "12. Репродуцибилност и верификација", page_break=True)
    add_body(doc, "Секој експеримент има единствен experiment ID и независен директориум под models_store/experiments. Manifest-от бележи timestamp, конфигурација, package versions, master seed, feature листа, sample counts, season completeness, data fingerprint и изворната политика за features. Ниту еден стар artifact не се презапишува. model_results.csv содржи резултат за task/context/model/fold; aggregate_results.csv ги содржи mean, standard deviation, rank и champion indicator; out_of_fold_predictions.csv.gz ги зачувува row-level предвидувањата; calibration/reliability и ablation резултатите имаат сопствени артефакти.")
    add_body(doc, "Репродуцибилниот финален команден запис е: docker compose run --rm ingestion python ml_pipeline/train_models.py --train-seasons 2021 2023 2024 --evaluation-seasons 2025 --min-train-seasons 2 --context all --experiment-id thesis-final-2025-holdout-20260620-r3 --seed 42 --artifact-output-dir models_store --model-output-dir models_store.", small=True)
    add_bullets(doc, [
        "unit тестови за temporal splitting, no-leakage гаранции, metric/calibration функции, baselines и artifact schema-и;",
        "API тестови со привремени artifact директориуми, latest-successful resolution и malformed-artifact случаи;",
        "frontend build/route верификација за Model Lab и prediction routes;",
        "рaчна проверка на клучни кориснички сценарија: сезонски преглед, анализа на трка, predictor, prediction comparison и Model Lab.",
    ])
    add_body(doc, "Репродуцибилноста не ја претвора малата сезонска база во голем примерок, но овозможува резултатите да се повторат, да се проверат и да се прошират кога ќе станат достапни нови комплетни сезони.")

    heading(doc, "13. Model Lab: презентација на експериментални докази")
    add_body(doc, "Model Lab ја трансформира перзистираната експериментална евиденција во читлив аналитички интерфејс. За избраниот experiment ID се прикажуваат dataset summary, temporal validation, held-out season, seed, feature context и ограничувања. Leaderboard-от може да се сортира, но champion ознаката произлегува од validation selection, а не од ретроспективно најдобар резултат на final holdout.")
    add_body(doc, "Графиконите се создаваат од запишани CSV артефакти, не од hard-coded вредности: leaderboard по задача, pre- против post-qualifying споредба, predicted-versus-actual/error графикони за регресија, ROC/precision-recall/calibration за класификација и ablation споредба. Ова е важна дистинкција: интерфејсот е приказ на репродуцибилен доказ, а не декоративен dashboard со непроверливи бројки.")
    add_body(doc, "За научното пишување се препорачуваат следните captions: „Model leaderboard за [задача], базиран на aggregate_results.csv од thesis-final-2025-holdout-20260620-r3“; „Споредба на pre-qualifying и post-qualifying champion-и на final holdout сезоната 2025“; и „Feature-ablation споредба на validation fold 2024; баровите не претставуваат каузален придонес на features“.")

    heading(doc, "14. Ограничувања", page_break=True)
    add_body(doc, "Резултатите треба да се толкуваат како аналитички прогнози, не како каузални објаснувања, betting совет или гаранција за исход. Safety car, red flag, судири, механички откажувања, казни и динамични стратегиски одлуки се важни, но не се познати на prediction cutoff. Исто така, временските услови можат да се сменат за време на трката; токму поради тоа target-race weather features се исклучени сè додека не постои timestamped forecast provenance.")
    add_bullets(doc, [
        "Недостасните резултати за 2022 ја намалуваат rolling-origin валидацијата на еден outer fold и ја зголемуваат неизвесноста на блиски model ranks.",
        "Podium класата е неурамнотежена; PR-AUC и калибрациските мерки се поинформативни од raw accuracy, но се чувствителни на мал број позитивни случаи.",
        "Регулативите, автомобилите и конкурентскиот поредок се менуваат меѓу сезони, па 2025 holdout не гарантира стабилност во идна година.",
        "Историскиот post-qualifying grid мора да остане усогласен со live feature contract-от; користење qualifying position како proxy треба да биде експлицитно документирано.",
        "Недостасни lap податоци, доцни замени на возачи, промени на тим и circuit-name matching можат да внесат measurement error или selection bias.",
    ])

    heading(doc, "15. Идни подобрувања", page_break=True)
    add_bullets(doc, [
        "Враќање или верификација на 2022 race-result coverage и додавање нови завршени сезони за повеќе outer temporal folds.",
        "Интеграција на timestamped weather-forecast извор со јасен pre-race cutoff пред повторно воведување на временски features.",
        "Проширување со race-cluster bootstrap интервали и подетална калибрациска анализа, особено за podium класата.",
        "Користење на официјален стартен grid за live inference или доследна qualifying-position proxy политика во train и inference.",
        "Посебен модел за DNF ризик и внимателно моделирање на pit-stop стратегија како дополнителни задачи, а не како скриени features.",
        "Per-driver explanation, monitoring на drift и верзионирани deployment одлуки во Model Lab.",
    ])
    add_body(doc, "Идните проширувања треба да ја задржат основната дисциплина: новите сигнали прво се проверуваат за availability и provenance, а новите модели се споредуваат на исти temporal folds и се промовираат само ако успешниот експеримент го оправдува тоа.")

    heading(doc, "16. Заклучок", page_break=True)
    add_body(doc, "ApexInsight демонстрира како една Formula 1 веб-платформа може да се префокусира во ML-ориентирана теза без да ги изгуби инженерските придобивки од базата, API-то и интерфејсот. Научниот центар на работата е репродуцибилниот експеримент за четири постојни targets и два информациски контексти. Строгото разделување на development од final holdout, исклучувањето на нецелосни 2026 податоци и зачувувањето на артефактите овозможуваат резултатите да се проверат и да не зависат од случајна UI демонстрација.")
    add_body(doc, "Финалната 2025 евалуација покажува дека post-qualifying информациите се поврзани со подобра изведба за завршна позиција, Top 10 и podium: MAE за завршна позиција е 3.258 наспроти 3.904 во pre-qualifying контекст, ROC-AUC за Top 10 е 0.837 наспроти 0.759, а PR-AUC за podium е 0.750 наспроти 0.523. Во исто време, резултатите не поддржуваат претерана генерализација: Ridge и Logistic Regression победуваат повеќе задачи, а pre-qualifying gain/loss не ја надминува zero-change базната линија. Таа комбинација од позитивни и негативни наоди е суштинска за коректна ML теза.")
    add_body(doc, "Платформата затоа треба да се разбере како аналитички систем со јасно документиран опфат. Нејзината вредност е во спојувањето на податоци, leakage-aware modelling, репродуцибилни артефакти и разбирливи визуелизации, при што ограничувањата — инциденти, стратегија, време, imbalance и сезонски промени — остануваат видливи наместо да бидат сокриени.")

    # Ensure core document metadata is useful without exposing personal content.
    props = doc.core_properties
    props.title = "ApexInsight: F1 ML Thesis Draft"
    props.subject = "Formula 1 prediction model comparison"
    props.author = "ApexInsight project team"
    props.comments = "Combined ML-centred thesis draft generated from local project documentation."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
